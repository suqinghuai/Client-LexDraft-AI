import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import configparser
import json
import os
import sys
import threading
import requests
from datetime import datetime

try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
    RESOURCE_DIR = sys._MEIPASS
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    RESOURCE_DIR = BASE_DIR

CONFIG_PATH = os.path.join(BASE_DIR, "config.ini")

COLORS = {
    "bg_dark": "#1B2A4A",
    "bg_sidebar": "#1E3258",
    "bg_sidebar_hover": "#264068",
    "bg_sidebar_active": "#2D4F7E",
    "accent_gold": "#C9A84C",
    "accent_gold_light": "#D4B96A",
    "bg_content": "#F5F6FA",
    "bg_card": "#FFFFFF",
    "bg_input": "#FAFBFD",
    "text_primary": "#1B2A4A",
    "text_secondary": "#5A6A85",
    "text_sidebar": "#C8D6E5",
    "text_sidebar_title": "#FFFFFF",
    "border_light": "#E2E8F0",
    "success_green": "#27AE60",
    "warning_orange": "#E67E22",
    "error_red": "#E74C3C",
    "btn_primary": "#2D4F7E",
    "btn_primary_hover": "#3A6199",
    "btn_success": "#27AE60",
    "btn_warning": "#E67E22",
    "btn_danger": "#E74C3C",
    "divider": "#2D4F7E",
}

FONT_FAMILY = "Microsoft YaHei UI"
FONT_FAMILY_MONO = "Consolas"

DOCUMENT_TREE = {
    "法院诉讼类": {
        "民事诉讼": {
            "一审": ["民事起诉状", "答辩状", "保全申请书"],
            "二审": ["民事上诉状"],
            "终审后": ["民事再审申请书"]
        },
        "行政诉讼": {
            "一审": ["行政起诉状"],
            "二审": ["行政上诉状"],
            "终审后": ["行政再审申请书"]
        },
        "刑事当事人可用文书": {
            "刑事自诉": ["刑事自诉状"],
            "刑事附带民事": ["刑事附带民事起诉状"],
            "刑事申诉": ["刑事申诉书"]
        }
    },
    "非诉讼诉求文书": {
        "行政机关举报": ["举报信"],
        "纪检监察专属": ["检举信（纪委监委）", "控告书（纪委监委）", "处分复查纪检申诉书"],
        "信访诉求材料": ["信访材料"],
        "仲裁调解保全执行": ["仲裁申请书", "调解申请书", "诉前保全申请书", "执行申请书"]
    }
}

def load_config():
    config = configparser.ConfigParser()
    config.read(CONFIG_PATH, encoding="utf-8")
    return {
        "url": config.get("AI", "URL", fallback=""),
        "api_key": config.get("AI", "API_KEY", fallback=""),
        "model": config.get("AI", "MODEL", fallback="gpt-3.5-turbo"),
        "temperature": config.getfloat("AI", "TEMPERATURE", fallback=0.7)
    }

def _find_prompt_file(doc_name):
    filename = f"{doc_name}.txt"
    user_path = os.path.join(BASE_DIR, "prompt", filename)
    if os.path.exists(user_path):
        return user_path
    builtin_path = os.path.join(RESOURCE_DIR, "prompt", filename)
    if os.path.exists(builtin_path):
        return builtin_path
    return None

def load_prompt(doc_name):
    prompt_file = _find_prompt_file(doc_name)
    if prompt_file:
        with open(prompt_file, "r", encoding="utf-8") as f:
            return f.read().strip()
    return ""

API_TIMEOUT = 120


def _do_api_call(config, messages):
    """统一 API 调用，供 call_api / call_api_chat 复用。"""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config['api_key']}"
    }
    payload = {
        "model": config["model"],
        "temperature": config["temperature"],
        "messages": messages
    }
    resp = None
    try:
        resp = requests.post(config["url"], headers=headers, json=payload, timeout=API_TIMEOUT)
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout:
        return "错误：请求超时，请检查网络连接或稍后重试。"
    except requests.exceptions.ConnectionError:
        return "错误：无法连接到API服务器，请检查网络和API地址。"
    except KeyError:
        resp_text = resp.text if resp is not None else "无响应"
        return f"错误：API返回格式异常。\n{resp_text}"
    except Exception as e:
        return f"错误：{str(e)}"


def call_api(config, system_prompt, user_input):
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_input}
    ]
    return _do_api_call(config, messages)


def call_api_chat(config, system_prompt, history_messages):
    messages = [{"role": "system", "content": system_prompt}] + history_messages
    return _do_api_call(config, messages)


# Word 导出不再依赖 xhtml2pdf / reportlab，直接用 python-docx，Word 自带中文字形。


class HoverButton(tk.Canvas):
    def __init__(self, parent, text, command=None, icon="", width=240,
                 bg=None, hover_bg=None, fg=None, hover_fg=None,
                 font_size=9, bold=False, active=False, indent=0, **kwargs):
        self._command = command
        self._bg = bg or COLORS["bg_sidebar"]
        self._hover_bg = hover_bg or COLORS["bg_sidebar_hover"]
        self._fg = fg or COLORS["text_sidebar"]
        self._hover_fg = hover_fg or "#FFFFFF"
        self._active = active
        self._current_bg = self._bg
        self._text = f"{icon}  {text}" if icon else text
        self._font_size = font_size
        self._bold = bold
        self._indent = indent

        super().__init__(parent, width=width, height=32, bg=self._bg,
                         highlightthickness=0, cursor="hand2", **kwargs)

        self._text_id = self.create_text(
            15 + indent, 16,
            text=self._text,
            fill=self._fg if not active else COLORS["accent_gold"],
            font=(FONT_FAMILY, font_size, "bold" if bold else "normal"),
            anchor="w"
        )

        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        self.bind("<ButtonPress-1>", self._on_click)

    def _on_enter(self, e):
        self._current_bg = self._hover_bg
        self.configure(bg=self._hover_bg)
        self.itemconfig(self._text_id, fill=self._hover_fg)

    def _on_leave(self, e):
        self._current_bg = self._bg
        self.configure(bg=self._bg)
        fill = COLORS["accent_gold"] if self._active else self._fg
        self.itemconfig(self._text_id, fill=fill)

    def _on_click(self, e):
        if self._command:
            self._command()

    def set_active(self, active):
        self._active = active
        fill = COLORS["accent_gold"] if active else self._fg
        self.itemconfig(self._text_id, fill=fill)


class CollapsibleSection(tk.Frame):
    def __init__(self, parent, title, level=1, expanded=True, **kwargs):
        super().__init__(parent, bg=COLORS["bg_dark"], **kwargs)
        self._expanded = expanded
        self._level = level
        self._title = title

        self._header = tk.Frame(self, bg=COLORS["bg_dark"], cursor="hand2")
        self._header.pack(fill=tk.X)

        indent = 10 if level == 1 else 20
        arrow_char = "▼" if expanded else "▶"
        font_size = 10 if level == 1 else 9
        fg_color = COLORS["accent_gold"] if level == 1 else COLORS["text_sidebar"]

        self._arrow = tk.Label(self._header, text=arrow_char,
                               font=(FONT_FAMILY, 8),
                               bg=COLORS["bg_dark"], fg=fg_color)
        self._arrow.pack(side=tk.LEFT, padx=(indent, 4), pady=(6 if level == 1 else 3, 2))

        self._title_label = tk.Label(self._header, text=title,
                                     font=(FONT_FAMILY, font_size, "bold"),
                                     bg=COLORS["bg_dark"], fg=fg_color,
                                     anchor="w")
        self._title_label.pack(side=tk.LEFT, pady=(6 if level == 1 else 3, 2))

        self._content = tk.Frame(self, bg=COLORS["bg_dark"])
        if self._expanded:
            self._content.pack(fill=tk.X)

        for widget in [self._header, self._arrow, self._title_label]:
            widget.bind("<ButtonPress-1>", self._toggle)
            widget.bind("<Enter>", self._on_header_enter)
            widget.bind("<Leave>", self._on_header_leave)

    def _toggle(self, event=None):
        if self._expanded:
            self._content.pack_forget()
            self._expanded = False
            self._arrow.configure(text="▶")
        else:
            self._content.pack(fill=tk.X)
            self._expanded = True
            self._arrow.configure(text="▼")

    def _on_header_enter(self, event=None):
        hover_bg = COLORS["bg_sidebar_hover"]
        self._header.configure(bg=hover_bg)
        self._arrow.configure(bg=hover_bg)
        self._title_label.configure(bg=hover_bg)

    def _on_header_leave(self, event=None):
        self._header.configure(bg=COLORS["bg_dark"])
        self._arrow.configure(bg=COLORS["bg_dark"])
        self._title_label.configure(bg=COLORS["bg_dark"])

    @property
    def content(self):
        return self._content

    @property
    def is_expanded(self):
        return self._expanded


class FlatButton(tk.Frame):
    def __init__(self, parent, text, command=None, color=None, hover_color=None,
                 fg="white", icon="", **kwargs):
        super().__init__(parent, **kwargs)
        self._command = command
        self._color = color or COLORS["btn_primary"]
        self._hover_color = hover_color or COLORS["btn_primary_hover"]
        self._fg = fg
        self._icon = icon

        self._label = tk.Label(
            self, text=f"{icon}  {text}" if icon else text,
            bg=self._color, fg=self._fg,
            font=(FONT_FAMILY, 10, "bold"),
            padx=16, pady=6, cursor="hand2"
        )
        self._label.pack(fill=tk.BOTH, expand=True)

        self._label.bind("<Enter>", lambda e: self._label.configure(bg=self._hover_color))
        self._label.bind("<Leave>", lambda e: self._label.configure(bg=self._color))
        self._label.bind("<ButtonPress-1>", lambda e: self._command() if self._command else None)

    def configure_state(self, state):
        if state == tk.DISABLED:
            self._label.configure(bg="#B0BEC5", cursor="arrow")
            self._color = "#B0BEC5"
            self._hover_color = "#B0BEC5"
        else:
            self._color = COLORS["btn_primary"]
            self._hover_color = COLORS["btn_primary_hover"]
            self._label.configure(bg=self._color, cursor="hand2")


class LegalDocumentApp:
    def __init__(self, root):
        self.root = root
        self.root.title("智能法律文书系统")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 650)
        self.root.configure(bg=COLORS["bg_content"])

        ico_path = os.path.join(BASE_DIR, "my.ico")
        if not os.path.exists(ico_path):
            ico_path = os.path.join(RESOURCE_DIR, "my.ico")
        if os.path.exists(ico_path):
            self.root.iconbitmap(ico_path)

        self.config = load_config()
        self.current_doc = None
        self.current_prompt = ""
        self._progress_running = False
        self._nav_buttons = []
        self.chat_messages = []          # 多轮对话历史
        self._chat_btn = None            # 智能法律顾问按钮（用于高亮态）
        self.legal_references_data = None          # 法条与类案数据
        self._last_refine_content = None           # 最近生成/优化的内容（供法条分析用）
        self._ref_panel = None                     # 右侧法条面板
        self._ref_panel_visible = False            # 面板展开状态
        self._ref_panel_content = None             # 面板内滚动区域Frame
        self._toggle_ref_btn = None                # 面板切换按钮

        self._build_ui()
        self._show_home()

        self.root.protocol("WM_DELETE_WINDOW", self._on_closing)

    def _on_closing(self):
        if messagebox.askyesno("确认退出", "确定要退出智能法律文书系统吗？"):
            self.root.destroy()

    def _build_ui(self):
        self.main_frame = tk.Frame(self.root, bg=COLORS["bg_content"])
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        self.nav_outer = tk.Frame(self.main_frame, bg=COLORS["bg_dark"], width=280)
        self.nav_outer.pack(side=tk.LEFT, fill=tk.Y)
        self.nav_outer.pack_propagate(False)

        self._build_nav()

        self.content_frame = tk.Frame(self.main_frame, bg=COLORS["bg_content"])
        self.content_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

    def _build_nav(self):
        header_frame = tk.Frame(self.nav_outer, bg=COLORS["bg_dark"], height=70)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)

        tk.Label(header_frame, text="⚖", font=(FONT_FAMILY, 22),
                 bg=COLORS["bg_dark"], fg=COLORS["accent_gold"]).pack(pady=(12, 0))
        tk.Label(header_frame, text="智能法律文书系统", font=(FONT_FAMILY, 11, "bold"),
                 bg=COLORS["bg_dark"], fg=COLORS["text_sidebar_title"]).pack()

        tk.Frame(self.nav_outer, bg=COLORS["divider"], height=1).pack(fill=tk.X, padx=15, pady=(5, 0))

        # ===== 新增：智能法律顾问入口（放在文书分类之前，更醒目） =====
        chat_top_frame = tk.Frame(self.nav_outer, bg=COLORS["bg_dark"])
        chat_top_frame.pack(fill=tk.X, padx=15, pady=(8, 4))
        self._chat_btn = HoverButton(
            chat_top_frame, text="智能法律顾问",
            icon="💬", command=self._show_chat,
            width=250, font_size=10, bold=True
        )
        self._chat_btn.pack(pady=2)

        tk.Label(self.nav_outer, text="文书分类", font=(FONT_FAMILY, 9),
                 bg=COLORS["bg_dark"], fg=COLORS["text_secondary"]).pack(anchor="w", padx=20, pady=(8, 3))

        canvas = tk.Canvas(self.nav_outer, bg=COLORS["bg_dark"], highlightthickness=0)
        scrollbar = tk.Scrollbar(self.nav_outer, orient=tk.VERTICAL, command=canvas.yview,
                                 bg=COLORS["bg_dark"], troughcolor=COLORS["bg_dark"])
        self.nav_inner = tk.Frame(canvas, bg=COLORS["bg_dark"])

        self.nav_inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.nav_inner, anchor="nw", width=260)
        canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind("<MouseWheel>", _on_mousewheel)
        self.nav_inner.bind("<MouseWheel>", _on_mousewheel)

        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self._nav_buttons = []
        for cat1_name, cat1_data in DOCUMENT_TREE.items():
            cat1_section = CollapsibleSection(self.nav_inner, cat1_name, level=1, expanded=True)
            cat1_section.pack(fill=tk.X, padx=2, pady=(4, 0))

            if isinstance(cat1_data, dict):
                for cat2_name, cat2_data in cat1_data.items():
                    if isinstance(cat2_data, dict):
                        cat2_section = CollapsibleSection(cat1_section.content, cat2_name, level=2, expanded=True)
                        cat2_section.pack(fill=tk.X, padx=2)

                        for cat3_name, docs in cat2_data.items():
                            for doc_name in docs:
                                btn = HoverButton(
                                    cat2_section.content, text=f"{cat3_name} · {doc_name}",
                                    command=lambda d=doc_name: self._select_doc(d),
                                    width=240, font_size=9, indent=10
                                )
                                btn.pack(padx=5, pady=1)
                                self._nav_buttons.append((btn, doc_name))
                    elif isinstance(cat2_data, list):
                        for doc_name in cat2_data:
                            btn = HoverButton(
                                cat1_section.content, text=f"{cat2_name} · {doc_name}",
                                command=lambda d=doc_name: self._select_doc(d),
                                width=240, font_size=9, indent=10
                            )
                            btn.pack(padx=5, pady=1)
                            self._nav_buttons.append((btn, doc_name))

        tk.Frame(self.nav_outer, bg=COLORS["divider"], height=1).pack(fill=tk.X, padx=15, pady=5)

        bottom_frame = tk.Frame(self.nav_outer, bg=COLORS["bg_dark"])
        bottom_frame.pack(fill=tk.X, padx=10, pady=(0, 10))

        home_btn = HoverButton(bottom_frame, text="首页", icon="🏠",
                               command=self._show_home, width=250, font_size=9, bold=True)
        home_btn.pack(pady=2)

        settings_btn = HoverButton(bottom_frame, text="系统设置", icon="⚙",
                                   command=self._show_settings, width=250, font_size=9, bold=True)
        settings_btn.pack(pady=2)

    def _clear_content(self):
        for w in self.content_frame.winfo_children():
            w.destroy()

    def _update_nav_active(self, doc_name=None):
        for btn, name in self._nav_buttons:
            btn.set_active(name == doc_name)
        # 智能法律顾问按钮的高亮态
        if self._chat_btn is not None:
            self._chat_btn.set_active(doc_name == "__chat__")

    def _show_chat(self):
        """智能法律顾问聊天界面。"""
        self._clear_content()
        self.current_doc = "__chat__"
        self._update_nav_active("__chat__")

        system_prompt = load_prompt("智能法律顾问")
        if not system_prompt:
            system_prompt = "你是一位资深、严谨的中国法律顾问，擅长用通俗易懂的语言解答法律问题，并在必要时提示用户寻求专业律师帮助。"
        self._chat_system_prompt = system_prompt

        header = tk.Frame(self.content_frame, bg=COLORS["bg_card"])
        header.pack(fill=tk.X, padx=20, pady=(15, 0))

        tk.Label(header, text="💬  智能法律顾问",
                 font=(FONT_FAMILY, 16, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=15, pady=12)
        tk.Label(header, text="✅ 已加载法律顾问提示词模板",
                 fg=COLORS["success_green"], bg=COLORS["bg_card"],
                 font=(FONT_FAMILY, 9)).pack(side=tk.LEFT, padx=10)

        # 对话历史显示区（上面大区域）
        chat_body = tk.Frame(self.content_frame, bg=COLORS["bg_content"])
        chat_body.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        chat_body.rowconfigure(0, weight=1)
        chat_body.columnconfigure(0, weight=1)

        history_card = tk.Frame(chat_body, bg=COLORS["bg_card"],
                                highlightbackground=COLORS["border_light"],
                                highlightthickness=1)
        history_card.grid(row=0, column=0, sticky="nsew", pady=(0, 5))

        history_header = tk.Frame(history_card, bg=COLORS["bg_card"])
        history_header.pack(fill=tk.X)
        tk.Label(history_header, text="  对话历史", font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=10, pady=8)
        tk.Frame(history_header, bg=COLORS["btn_primary"], height=2).pack(fill=tk.X, padx=10)

        self.chat_history = scrolledtext.ScrolledText(
            history_card, wrap=tk.WORD,
            font=(FONT_FAMILY, 10),
            bg=COLORS["bg_input"], fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"],
            relief=tk.FLAT, borderwidth=0,
            selectbackground=COLORS["btn_primary"],
            state=tk.DISABLED
        )
        self.chat_history.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        # 自定义两种文字样式：我 / 法律顾问
        self.chat_history.tag_configure("me", foreground=COLORS["btn_primary"],
                                        font=(FONT_FAMILY, 10, "bold"))
        self.chat_history.tag_configure("ai", foreground=COLORS["accent_gold"],
                                        font=(FONT_FAMILY, 10, "bold"))
        self.chat_history.tag_configure("text", foreground=COLORS["text_primary"],
                                        font=(FONT_FAMILY, 10))

        # 欢迎语
        self._append_chat_text("法律顾问", "您好，我是您的智能法律顾问。您可以向我咨询任何法律问题，我会尽我所能为您解答。\n请注意：本顾问的回答仅供参考，不构成正式法律意见。", role="ai")

        # 工具栏（下面的输入区）
        toolbar = tk.Frame(chat_body, bg=COLORS["bg_content"])
        toolbar.grid(row=1, column=0, sticky="ew", pady=5)

        self.chat_send_btn = FlatButton(toolbar, text="发送", icon="➤",
                                        command=self._send_chat_message,
                                        color=COLORS["btn_primary"],
                                        hover_color=COLORS["btn_primary_hover"])
        self.chat_send_btn.pack(side=tk.LEFT, padx=(0, 8))

        self.chat_clear_btn = FlatButton(toolbar, text="清空对话", icon="🧹",
                                         command=self._clear_chat,
                                         color="#7F8C8D", hover_color="#95A5A6")
        self.chat_clear_btn.pack(side=tk.LEFT, padx=(0, 8))

        self.chat_status_var = tk.StringVar(value="就绪")
        self.chat_status_label = tk.Label(toolbar, textvariable=self.chat_status_var,
                                          font=(FONT_FAMILY, 9),
                                          bg=COLORS["bg_content"], fg=COLORS["text_secondary"])
        self.chat_status_label.pack(side=tk.RIGHT, padx=10)

        # 输入区（最下面，多行）
        input_card = tk.Frame(chat_body, bg=COLORS["bg_card"],
                              highlightbackground=COLORS["border_light"],
                              highlightthickness=1)
        input_card.grid(row=2, column=0, sticky="ew", pady=(5, 0))
        input_card.rowconfigure(1, weight=1)
        input_card.columnconfigure(0, weight=1)

        input_header = tk.Frame(input_card, bg=COLORS["bg_card"])
        input_header.grid(row=0, column=0, sticky="ew")
        tk.Label(input_header, text="  输入您的问题（Ctrl+Enter 发送）", font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=10, pady=8)
        tk.Frame(input_header, bg=COLORS["accent_gold"], height=2).pack(fill=tk.X, padx=10)

        self.chat_input = scrolledtext.ScrolledText(
            input_card, wrap=tk.WORD,
            font=(FONT_FAMILY, 10), height=4,
            bg=COLORS["bg_input"], fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"],
            relief=tk.FLAT, borderwidth=0,
            selectbackground=COLORS["btn_primary"]
        )
        self.chat_input.grid(row=1, column=0, sticky="ew", padx=15, pady=10)
        self.chat_input.bind("<Control-Return>", lambda e: self._send_chat_message())

    def _append_chat_text(self, title, content, role="user"):
        """向对话历史区追加一条消息。"""
        self.chat_history.config(state=tk.NORMAL)
        self.chat_history.insert(tk.END, f"【{title}】\n", "ai" if role == "ai" else "me")
        self.chat_history.insert(tk.END, f"{content}\n\n", "text")
        self.chat_history.see(tk.END)
        self.chat_history.config(state=tk.DISABLED)

    def _send_chat_message(self):
        user_text = self.chat_input.get("1.0", tk.END).strip()
        if not user_text:
            return
        if not self.config["api_key"] or self.config["api_key"] == "sk-xxxxxxxx":
            messagebox.showwarning("提示", "请先在「设置」中配置有效的API Key。")
            return

        # 显示我的问题并清空输入框
        self._append_chat_text("我", user_text, role="me")
        self.chat_input.delete("1.0", tk.END)

        self.chat_send_btn.configure_state(tk.DISABLED)
        self.chat_status_var.set("⏳ 顾问正在思考...")
        self.chat_status_label.configure(fg=COLORS["btn_primary"])

        # 加入历史
        self.chat_messages.append({"role": "user", "content": user_text})

        system_prompt = getattr(self, "_chat_system_prompt", "")

        def worker():
            reply = call_api_chat(self.config, system_prompt, self.chat_messages)
            self.root.after(0, self._on_chat_reply, reply)

        threading.Thread(target=worker, daemon=True).start()

    def _on_chat_reply(self, reply):
        self.chat_messages.append({"role": "assistant", "content": reply})
        self._append_chat_text("法律顾问", reply, role="ai")
        self.chat_send_btn.configure_state(tk.NORMAL)
        self.chat_status_var.set("✅ 完成")
        self.chat_status_label.configure(fg=COLORS["success_green"])

    def _clear_chat(self):
        self.chat_messages = []
        self.chat_history.config(state=tk.NORMAL)
        self.chat_history.delete("1.0", tk.END)
        self.chat_history.config(state=tk.DISABLED)
        self.chat_status_var.set("对话已清空")
        self.chat_status_label.configure(fg=COLORS["text_secondary"])

    def _show_home(self):
        self._clear_content()
        self.current_doc = None
        self._update_nav_active(None)

        home_frame = tk.Frame(self.content_frame, bg=COLORS["bg_content"])
        home_frame.pack(fill=tk.BOTH, expand=True)

        hero_frame = tk.Frame(home_frame, bg=COLORS["bg_card"])
        hero_frame.pack(fill=tk.X, padx=30, pady=(30, 0))

        tk.Label(hero_frame, text="⚖", font=(FONT_FAMILY, 36),
                 bg=COLORS["bg_card"], fg=COLORS["accent_gold"]).pack(pady=(30, 0))
        tk.Label(hero_frame, text="智能法律文书系统",
                 font=(FONT_FAMILY, 24, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(pady=(5, 0))
        tk.Label(hero_frame, text="基于大模型的司法文书草拟 · 修正 · 优化系统",
                 font=(FONT_FAMILY, 11),
                 bg=COLORS["bg_card"], fg=COLORS["text_secondary"]).pack(pady=(0, 25))

        cards_frame = tk.Frame(home_frame, bg=COLORS["bg_content"])
        cards_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)

        cards_frame.columnconfigure(0, weight=1)
        cards_frame.columnconfigure(1, weight=1)
        cards_frame.rowconfigure(0, weight=1)
        cards_frame.rowconfigure(1, weight=1)

        card_data = [
            ("📋", "选择文书", "从左侧菜单选择\n需要草拟的文书类型", COLORS["btn_primary"]),
            ("✍️", "输入材料", "填写案件材料\n和相关信息", COLORS["btn_success"]),
            ("🤖", "AI 生成", "调用大模型生成\n规范法律文书", COLORS["accent_gold"]),
            ("📄", "导出结果", "查看、复制或\n导出文书文件", COLORS["btn_warning"]),
        ]

        for i, (icon, title, desc, color) in enumerate(card_data):
            row, col = divmod(i, 2)
            card = tk.Frame(cards_frame, bg=COLORS["bg_card"],
                            highlightbackground=COLORS["border_light"],
                            highlightthickness=1)
            card.grid(row=row, column=col, padx=8, pady=8, sticky="nsew")

            tk.Label(card, text=icon, font=(FONT_FAMILY, 28),
                     bg=COLORS["bg_card"]).pack(pady=(15, 5))
            tk.Label(card, text=title, font=(FONT_FAMILY, 13, "bold"),
                     bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack()
            tk.Frame(card, bg=color, height=3, width=40).pack(pady=5)
            tk.Label(card, text=desc, font=(FONT_FAMILY, 9),
                     bg=COLORS["bg_card"], fg=COLORS["text_secondary"],
                     justify=tk.CENTER).pack(pady=(0, 15))

        total = 0
        for cat1_data in DOCUMENT_TREE.values():
            if isinstance(cat1_data, dict):
                for cat2_data in cat1_data.values():
                    if isinstance(cat2_data, dict):
                        for docs in cat2_data.values():
                            total += len(docs)
                    else:
                        total += len(cat2_data)

        stat_bar = tk.Frame(home_frame, bg=COLORS["bg_dark"], height=40)
        stat_bar.pack(fill=tk.X, padx=30, pady=(0, 20))
        stat_bar.pack_propagate(False)
        tk.Label(stat_bar, text=f"  当前支持 {total} 种法律文书类型  ·  法院诉讼类 + 非诉讼诉求文书",
                 font=(FONT_FAMILY, 9), bg=COLORS["bg_dark"],
                 fg=COLORS["text_sidebar"]).pack(side=tk.LEFT, padx=10, pady=8)

    def _select_doc(self, doc_name):
        self.current_doc = doc_name
        self.current_prompt = load_prompt(doc_name)
        self._update_nav_active(doc_name)
        self._show_editor()

    def _show_editor(self):
        self._clear_content()

        header = tk.Frame(self.content_frame, bg=COLORS["bg_card"])
        header.pack(fill=tk.X, padx=20, pady=(15, 0))

        tk.Label(header, text=f"📝  {self.current_doc}",
                 font=(FONT_FAMILY, 16, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=15, pady=12)

        if self.current_prompt:
            tk.Label(header, text="✅ 已加载专业提示词模板",
                     fg=COLORS["success_green"], bg=COLORS["bg_card"],
                     font=(FONT_FAMILY, 9)).pack(side=tk.LEFT, padx=10)
        else:
            tk.Label(header, text="⚠ 通用模式（未找到专用模板）",
                     fg=COLORS["warning_orange"], bg=COLORS["bg_card"],
                     font=(FONT_FAMILY, 9)).pack(side=tk.LEFT, padx=10)

        # ===== 主体区域：左右分栏 =====
        body_container = tk.Frame(self.content_frame, bg=COLORS["bg_content"])
        body_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        body_container.columnconfigure(0, weight=1)
        body_container.columnconfigure(1, weight=0)
        body_container.columnconfigure(2, weight=0)

        # ---- 左侧：编辑器 ----
        body = tk.Frame(body_container, bg=COLORS["bg_content"])
        body.grid(row=0, column=0, sticky="nsew")
        body.rowconfigure(0, weight=1)
        body.rowconfigure(2, weight=1)
        body.columnconfigure(0, weight=1)

        input_card = tk.Frame(body, bg=COLORS["bg_card"],
                              highlightbackground=COLORS["border_light"],
                              highlightthickness=1)
        input_card.grid(row=0, column=0, sticky="nsew", pady=(0, 5))

        input_header = tk.Frame(input_card, bg=COLORS["bg_card"])
        input_header.pack(fill=tk.X)
        tk.Label(input_header, text="  输入材料与信息", font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=10, pady=8)
        tk.Frame(input_header, bg=COLORS["btn_primary"], height=2).pack(fill=tk.X, padx=10)

        self.input_text = scrolledtext.ScrolledText(
            input_card, wrap=tk.WORD,
            font=(FONT_FAMILY, 10), height=8,
            bg=COLORS["bg_input"], fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"],
            relief=tk.FLAT, borderwidth=0,
            selectbackground=COLORS["btn_primary"]
        )
        self.input_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        placeholder = "请在此输入您的案件材料和信息，例如：\n\n当事人信息：原告张三，被告李四\n案件事实：...\n诉讼请求：...\n"
        self.input_text.insert("1.0", placeholder)
        self.input_text.bind("<FocusIn>", self._clear_placeholder)

        toolbar = tk.Frame(body, bg=COLORS["bg_content"])
        toolbar.grid(row=1, column=0, sticky="ew", pady=5)

        self.gen_btn = FlatButton(toolbar, text="生成文书", icon="🚀",
                                  command=self._generate,
                                  color=COLORS["btn_primary"],
                                  hover_color=COLORS["btn_primary_hover"])
        self.gen_btn.pack(side=tk.LEFT, padx=(0, 8))

        self.refine_btn = FlatButton(toolbar, text="修正优化", icon="🔧",
                                     command=self._refine,
                                     color=COLORS["btn_success"],
                                     hover_color="#2ECC71")
        self.refine_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.refine_btn.configure_state(tk.DISABLED)

        self.export_btn = FlatButton(toolbar, text="导出文件", icon="💾",
                                     command=self._export,
                                     color=COLORS["btn_warning"],
                                     hover_color="#F39C12")
        self.export_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.export_btn.configure_state(tk.DISABLED)

        self.copy_btn = FlatButton(toolbar, text="复制结果", icon="📋",
                                   command=self._copy_result,
                                   color="#7F8C8D", hover_color="#95A5A6")
        self.copy_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.copy_btn.configure_state(tk.DISABLED)

        # ===== 右侧面板切换按钮（替代原来的法条类案按钮） =====
        self._toggle_ref_btn = FlatButton(toolbar, text="📁 法条", icon="",
                                          command=self._toggle_ref_panel,
                                          color="#8E44AD", hover_color="#9B59B6")
        self._toggle_ref_btn.pack(side=tk.LEFT, padx=(0, 8))
        self._toggle_ref_btn.configure_state(tk.DISABLED)

        self.status_var = tk.StringVar(value="就绪")
        self.status_label = tk.Label(toolbar, textvariable=self.status_var,
                                     font=(FONT_FAMILY, 9),
                                     bg=COLORS["bg_content"], fg=COLORS["text_secondary"])
        self.status_label.pack(side=tk.RIGHT, padx=10)

        self.progress = ttk.Progressbar(toolbar, mode="indeterminate", length=120)

        output_card = tk.Frame(body, bg=COLORS["bg_card"],
                               highlightbackground=COLORS["border_light"],
                               highlightthickness=1)
        output_card.grid(row=2, column=0, sticky="nsew", pady=(5, 0))

        output_header = tk.Frame(output_card, bg=COLORS["bg_card"])
        output_header.pack(fill=tk.X)
        tk.Label(output_header, text="  生成结果", font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=10, pady=8)
        tk.Frame(output_header, bg=COLORS["accent_gold"], height=2).pack(fill=tk.X, padx=10)

        self.output_text = scrolledtext.ScrolledText(
            output_card, wrap=tk.WORD,
            font=(FONT_FAMILY, 10), height=8,
            bg=COLORS["bg_input"], fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"],
            relief=tk.FLAT, borderwidth=0,
            selectbackground=COLORS["btn_primary"],
            state=tk.DISABLED
        )
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        # ---- 分隔线 ----
        tk.Frame(body_container, bg=COLORS["border_light"], width=1).grid(row=0, column=1, sticky="ns", padx=5)

        # ---- 右侧：法条与类案可折叠面板 ----
        self._build_ref_panel_ui(body_container)
        self._ref_panel_visible = False
        self._ref_panel.grid_remove()

    def _clear_placeholder(self, event):
        placeholder = "请在此输入您的案件材料和信息，例如：\n\n当事人信息：原告张三，被告李四\n案件事实：...\n诉讼请求：...\n"
        if self.input_text.get("1.0", tk.END).strip() == placeholder.strip():
            self.input_text.delete("1.0", tk.END)

    def _start_progress(self):
        self._progress_running = True
        self.progress.pack(side=tk.RIGHT, padx=5)
        self.progress.start(15)

    def _stop_progress(self):
        self._progress_running = False
        self.progress.stop()
        self.progress.pack_forget()

    def _generate(self):
        user_input = self.input_text.get("1.0", tk.END).strip()
        if not user_input:
            messagebox.showwarning("提示", "请先输入案件材料和信息。")
            return

        if not self.config["api_key"] or self.config["api_key"] == "sk-xxxxxxxx":
            messagebox.showwarning("提示", "请先在「设置」中配置有效的API Key。")
            return

        self.gen_btn.configure_state(tk.DISABLED)
        self.refine_btn.configure_state(tk.DISABLED)
        self._start_progress()
        self.status_var.set("⏳ 正在生成文书...")
        self.status_label.configure(fg=COLORS["btn_primary"])

        system_prompt = self.current_prompt if self.current_prompt else (
            f"你是一位专业的中国法律文书撰写专家。请根据用户提供的材料和信息，"
            f"草拟一份规范的「{self.current_doc}」。要求格式规范、用语专业、逻辑清晰。"
        )

        full_user_input = f"请根据以下材料草拟「{self.current_doc}」：\n\n{user_input}"

        def worker():
            result = call_api(self.config, system_prompt, full_user_input)
            self.root.after(0, self._on_result, result)

        threading.Thread(target=worker, daemon=True).start()

    def _refine(self):
        current_result = self.output_text.get("1.0", tk.END).strip()
        if not current_result:
            current_result = self.input_text.get("1.0", tk.END).strip()
            placeholder = "请在此输入您的案件材料和信息，例如：\n\n当事人信息：原告张三，被告李四\n案件事实：...\n诉讼请求：...\n"
            if not current_result or current_result == placeholder.strip():
                messagebox.showwarning("提示", "没有可优化的文书内容，请先生成文书或在输入区粘贴文书。")
                return

        self.refine_btn.configure_state(tk.DISABLED)
        self.gen_btn.configure_state(tk.DISABLED)
        self._start_progress()
        self.status_var.set("⏳ 正在修正优化...")
        self.status_label.configure(fg=COLORS["btn_success"])

        system_prompt = load_prompt("修改优化")
        if not system_prompt:
            system_prompt = (
                "你是一位精通中国法律的专业律师和法律文书撰写优化专家，擅长草拟规范的法律文书。"
                "请根据用户提供的文案，对其进行修改优化，确保用语严谨、格式规范。"
                "直接返回优化的最终结果和指出优化的地方。"
            )

        user_input = f"请对以下文书进行修改优化：\n\n{current_result}"

        def worker():
            result = call_api(self.config, system_prompt, user_input)
            self.root.after(0, self._on_result, result)

        threading.Thread(target=worker, daemon=True).start()

    def _on_result(self, result):
        self._stop_progress()
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert("1.0", result)
        self.output_text.config(state=tk.DISABLED)

        self.gen_btn.configure_state(tk.NORMAL)
        self.refine_btn.configure_state(tk.NORMAL)
        self.export_btn.configure_state(tk.NORMAL)
        self.copy_btn.configure_state(tk.NORMAL)
        self.status_var.set("✅ 完成")
        self.status_label.configure(fg=COLORS["success_green"])

        # 保存最近生成的文书内容，自动触发法条与类案检索
        self._last_refine_content = result
        self.legal_references_data = None
        self._toggle_ref_btn.configure_state(tk.DISABLED)
        threading.Thread(target=self._fetch_legal_references, args=(result,), daemon=True).start()

    def _fetch_legal_references(self, content):
        """异步获取法条分析与类案推送。"""
        if not self.config["api_key"] or self.config["api_key"] == "sk-xxxxxxxx":
            self.root.after(0, lambda: self.status_var.set("⚠ 未配置API，跳过法条检索"))
            return

        system_prompt = (
            "你是一位精通中国法律的高级法律检索专家。请根据以下法律文书内容，"
            "进行专业的法律检索分析。\n\n"
            "请分析以下三个方面，必须以 **JSON 数组格式** 返回（不要用markdown代码块包裹，纯JSON文本）：\n\n"
            "1. legal_provisions（核心法律条文）：列出最相关的法律条文，每项包含：\n"
            "   - law: 法律名称（如《中华人民共和国民法典》）\n"
            "   - article: 具体条、款、项（如第一千二百一十七条）\n"
            "   - summary: 该条文核心内容摘要（20字以内）\n"
            "   - relevance: 与该案的相关性说明\n\n"
            "2. reference_cases（相似参考案例）：推荐2-3个类似案例，每项包含：\n"
            "   - type: 案例类型/案由\n"
            "   - summary: 裁判要旨（50字以内）\n"
            "   - value: 参考价值说明\n\n"
            "3. risk_warnings（诉讼风险提示）：基于文书识别风险，每项包含：\n"
            "   - type: 风险类型（证据风险/程序风险/实体风险/其他）\n"
            "   - content: 具体风险描述与建议\n\n"
            "返回格式示例（纯JSON，不要任何其他文字）：\n"
            '{"legal_provisions":[{"law":"《中华人民共和国民法典》","article":"第一千二百一十七条","summary":"好意同乘责任减轻","relevance":"本案属于好意同乘情形"}],'
            '"reference_cases":[{"type":"机动车交通事故责任纠纷","summary":"驾驶人无证驾驶，同乘人知情仍搭乘，减轻驾驶人责任","value":"与本案事实高度相似，具有重要参考价值"}],'
            '"risk_warnings":[{"type":"证据风险","content":"建议收集行车记录仪视频以证明同乘人知情"}]}'
        )
        user_input = f"请对以下文书进行法条检索与类案分析：\n\n{content[:4000]}"

        try:
            result = call_api(self.config, system_prompt, user_input)
            self.root.after(0, self._on_references_result, result)
        except Exception as e:
            self.root.after(0, lambda: self.status_var.set(f"⚠ 法条检索失败: {str(e)}"))

    def _on_references_result(self, raw_result):
        """处理法条类案API返回结果。"""
        json_str = raw_result.strip()
        if json_str.startswith("```"):
            start = json_str.find("\n")
            end = json_str.rfind("```")
            if start != -1 and end != -1:
                json_str = json_str[start:end].strip()
        brace_start = json_str.find("{")
        brace_end = json_str.rfind("}")
        if brace_start != -1 and brace_end != -1:
            json_str = json_str[brace_start:brace_end + 1]

        try:
            data = json.loads(json_str)
            self.legal_references_data = {
                "legal_provisions": data.get("legal_provisions", []),
                "reference_cases": data.get("reference_cases", []),
                "risk_warnings": data.get("risk_warnings", []),
            }
        except json.JSONDecodeError:
            self.legal_references_data = {
                "_raw": raw_result,
                "legal_provisions": [],
                "reference_cases": [],
                "risk_warnings": [],
            }
        # 数据就绪后自动填充面板
        self.root.after(0, self._on_references_ready)

    def _on_references_ready(self):
        """法条数据就绪：填充右侧面板并自动展开。"""
        self._populate_ref_panel()
        if self._ref_panel:
            if not self._ref_panel_visible:
                self._toggle_ref_panel()
        self._toggle_ref_btn.configure_state(tk.NORMAL)
        data = self.legal_references_data or {}
        n_law = len(data.get("legal_provisions", []))
        n_case = len(data.get("reference_cases", []))
        n_risk = len(data.get("risk_warnings", []))
        self.status_var.set(f"✅ 完成  📚 法条{n_law}条 · 类案{n_case}件 · 风险{n_risk}项")
        self.status_label.configure(fg=COLORS["success_green"])

    def _build_ref_panel_ui(self, parent):
        """构建右侧可折叠法条面板的UI结构。"""
        self._ref_panel = tk.Frame(parent, bg=COLORS["bg_card"],
                                    highlightbackground=COLORS["border_light"],
                                    highlightthickness=1)
        self._ref_panel.grid(row=0, column=2, sticky="nsew")
        self._ref_panel.grid_propagate(False)

        # 面板标题栏（点击可折叠）
        panel_header = tk.Frame(self._ref_panel, bg=COLORS["bg_card"], height=36)
        panel_header.pack(fill=tk.X)
        panel_header.pack_propagate(False)

        tk.Label(panel_header, text="  📚  法条与类案",
                 font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=8, pady=6)
        self._panel_close_btn = tk.Label(panel_header, text="  ✕  ",
                                          font=(FONT_FAMILY, 10),
                                          bg=COLORS["bg_card"], fg=COLORS["text_secondary"],
                                          cursor="hand2")
        self._panel_close_btn.pack(side=tk.RIGHT, padx=5)
        self._panel_close_btn.bind("<ButtonPress-1>", lambda e: self._toggle_ref_panel())

        tk.Frame(panel_header, bg=COLORS["accent_gold"], height=2).pack(fill=tk.X, padx=8)

        # 滚动内容区域
        panel_body = tk.Frame(self._ref_panel, bg=COLORS["bg_card"])
        panel_body.pack(fill=tk.BOTH, expand=True)
        panel_body.rowconfigure(0, weight=1)
        panel_body.columnconfigure(0, weight=1)

        canvas = tk.Canvas(panel_body, bg=COLORS["bg_card"], highlightthickness=0)
        scrollbar = tk.Scrollbar(panel_body, orient=tk.VERTICAL, command=canvas.yview,
                                 bg=COLORS["bg_card"], troughcolor=COLORS["bg_card"])
        self._ref_panel_content = tk.Frame(canvas, bg=COLORS["bg_card"])

        self._ref_panel_content.bind("<Configure>",
                                      lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self._ref_panel_content, anchor="nw",
                              width=300)
        canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        def _bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", _on_mousewheel)

        def _unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")

        canvas.bind("<Enter>", _bind_mousewheel)
        canvas.bind("<Leave>", _unbind_mousewheel)

        # 初始化占位提示
        self._ref_placeholder = tk.Label(
            self._ref_panel_content,
            text="\n\n生成文书后，\n法条与类案将\n自动显示在这里\n\n📚",
            font=(FONT_FAMILY, 10),
            bg=COLORS["bg_card"], fg=COLORS["text_secondary"],
            justify=tk.CENTER
        )
        self._ref_placeholder.pack(expand=True, pady=40)

        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

    def _toggle_ref_panel(self):
        """切换右侧面板的展开/折叠状态。"""
        if not self._ref_panel:
            return
        if self._ref_panel_visible:
            self._ref_panel.grid_remove()
            self._ref_panel_visible = False
            if self._toggle_ref_btn:
                self._toggle_ref_btn._label.config(text="📁 法条")
            self.status_var.set("📁 法条面板已关闭")
            self.status_label.configure(fg=COLORS["text_secondary"])
        else:
            self._ref_panel.grid()
            self._ref_panel_visible = True
            if self._toggle_ref_btn:
                self._toggle_ref_btn._label.config(text="📂 法条")
            self.status_var.set("📂 法条面板已展开")
            self.status_label.configure(fg=COLORS["success_green"])

    def _populate_ref_panel(self):
        """用法条类案数据填充右侧面板。"""
        # 清空占位
        for w in self._ref_panel_content.winfo_children():
            w.destroy()

        data = self.legal_references_data
        if not data:
            tk.Label(self._ref_panel_content,
                     text="暂无数据", font=(FONT_FAMILY, 9),
                     bg=COLORS["bg_card"], fg=COLORS["text_secondary"]
                     ).pack(pady=20)
            return

        # 如果JSON解析失败，显示原始数据
        if "_raw" in data:
            tk.Label(self._ref_panel_content,
                     text="⚠ JSON解析失败，显示原始返回：",
                     font=(FONT_FAMILY, 9, "bold"),
                     bg=COLORS["bg_card"], fg=COLORS["error_red"]
                     ).pack(anchor="w", padx=10, pady=5)
            raw_text = tk.Text(self._ref_panel_content,
                               font=(FONT_FAMILY_MONO, 8), height=15,
                               bg=COLORS["bg_input"], fg=COLORS["text_primary"],
                               relief=tk.FLAT, borderwidth=0)
            raw_text.pack(fill=tk.BOTH, expand=True, padx=8, pady=5)
            raw_text.insert("1.0", data["_raw"])
            raw_text.config(state=tk.DISABLED)
            return

        parent = self._ref_panel_content
        pad = 8

        # ----- 1. 核心法律条文 -----
        self._add_panel_section(
            parent, "⚖  核心法律条文", COLORS["btn_primary"],
            data.get("legal_provisions", []),
            lambda item: (f"📜 {item.get('law', '')}\n"
                          f"   条款：{item.get('article', '')}\n"
                          f"   {item.get('summary', '')}\n"
                          f"   {item.get('relevance', '')}")
        )

        # ----- 2. 相似参考案例 -----
        self._add_panel_section(
            parent, "📋  相似参考案例", COLORS["success_green"],
            data.get("reference_cases", []),
            lambda item: (f"🏛 {item.get('type', '')}\n"
                          f"   {item.get('summary', '')}\n"
                          f"   💡 {item.get('value', '')}")
        )

        # ----- 3. 诉讼风险提示 -----
        risk_items = data.get("risk_warnings", [])
        risk_frame = tk.Frame(parent, bg=COLORS["bg_card"])
        risk_frame.pack(fill=tk.X, padx=pad, pady=2)

        tk.Label(risk_frame, text="⚠  诉讼风险提示",
                 font=(FONT_FAMILY, 9, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["error_red"]).pack(anchor="w")
        tk.Frame(risk_frame, bg=COLORS["error_red"], height=1).pack(fill=tk.X, pady=2)

        if not risk_items:
            tk.Label(risk_frame, text="暂无明显的诉讼风险",
                     font=(FONT_FAMILY, 8),
                     bg=COLORS["bg_card"], fg=COLORS["text_secondary"]
                     ).pack(anchor="w", pady=4)
        else:
            for item in risk_items:
                rtype = item.get("type", "其他")
                rcontent = item.get("content", "")
                type_colors = {
                    "证据风险": COLORS["warning_orange"],
                    "程序风险": COLORS["btn_primary"],
                    "实体风险": COLORS["error_red"],
                }
                color = type_colors.get(rtype, COLORS["warning_orange"])
                item_f = tk.Frame(risk_frame, bg=COLORS["bg_card"])
                item_f.pack(fill=tk.X, pady=2)
                tk.Label(item_f, text=f"⚠ [{rtype}]",
                         font=(FONT_FAMILY, 8, "bold"),
                         bg=COLORS["bg_card"], fg=color).pack(anchor="w")
                tk.Label(item_f, text=rcontent,
                         font=(FONT_FAMILY, 8),
                         bg=COLORS["bg_card"], fg=COLORS["text_primary"],
                         wraplength=280, justify=tk.LEFT).pack(anchor="w", padx=4)

        # ----- 4. 底部操作按钮 -----
        btn_frame = tk.Frame(parent, bg=COLORS["bg_card"])
        btn_frame.pack(fill=tk.X, padx=pad, pady=8)
        refresh_btn = tk.Label(btn_frame, text="🔄 重新检索",
                                font=(FONT_FAMILY, 8),
                                bg=COLORS["bg_card"], fg=COLORS["btn_primary"],
                                cursor="hand2")
        refresh_btn.pack()
        refresh_btn.bind("<ButtonPress-1>", lambda e: self._re_fetch_references())

    def _add_panel_section(self, parent, title, accent_color, items, format_fn):
        """在面板中添加一个分类区块（法条/案例）。"""
        section = tk.Frame(parent, bg=COLORS["bg_card"])
        section.pack(fill=tk.X, padx=8, pady=2)

        tk.Label(section, text=title,
                 font=(FONT_FAMILY, 9, "bold"),
                 bg=COLORS["bg_card"], fg=accent_color).pack(anchor="w")
        tk.Label(section, text=f"共 {len(items)} 条",
                 font=(FONT_FAMILY, 7),
                 bg=COLORS["bg_card"], fg=COLORS["text_secondary"]
                 ).pack(anchor="w", padx=4)
        tk.Frame(section, bg=accent_color, height=1).pack(fill=tk.X, pady=2)

        if not items:
            tk.Label(section, text="暂无相关数据",
                     font=(FONT_FAMILY, 8),
                     bg=COLORS["bg_card"], fg=COLORS["text_secondary"]
                     ).pack(anchor="w", pady=4)
        else:
            for i, item in enumerate(items, 1):
                item_f = tk.Frame(section, bg=COLORS["bg_card"])
                item_f.pack(fill=tk.X, pady=2)

                text = f"{i}. {format_fn(item)}"
                tk.Label(item_f, text=text,
                         font=(FONT_FAMILY, 8),
                         bg=COLORS["bg_card"], fg=COLORS["text_primary"],
                         wraplength=280, justify=tk.LEFT, anchor="w"
                         ).pack(fill=tk.X)

                if i < len(items):
                    tk.Frame(item_f, bg=COLORS["border_light"], height=1).pack(fill=tk.X, pady=1)

    def _re_fetch_references(self):
        """手动重新检索法条类案。"""
        if not self._last_refine_content:
            messagebox.showwarning("提示", "没有已生成的文书内容可供检索。")
            return
        self._toggle_ref_btn.configure_state(tk.DISABLED)
        self.status_var.set("⏳ 正在重新检索法条与类案...")
        self.status_label.configure(fg=COLORS["btn_primary"])
        # 面板显示加载中
        if self._ref_panel_content:
            for w in self._ref_panel_content.winfo_children():
                w.destroy()
            tk.Label(self._ref_panel_content,
                     text="\n⏳ 正在检索...\n\n请稍候",
                     font=(FONT_FAMILY, 10),
                     bg=COLORS["bg_card"], fg=COLORS["text_secondary"],
                     justify=tk.CENTER).pack(expand=True, pady=40)
        threading.Thread(target=self._fetch_legal_references,
                         args=(self._last_refine_content,), daemon=True).start()

    def _markdown_to_docx(self, md_text, output_path):
        """把 markdown 文本转换成 Word(.docx)。

        思路：markdown -> markdown 转 HTML 后解析，再遍历 DOM 转为 docx 段落/表格。
        使用 python-docx 渲染 Word 本身自带中文字形，不会出现黑色方块。
        """
        html_body = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
        soup = BeautifulSoup('<html><body>' + html_body + '</body></html>', 'html.parser')

        doc = Document()

        # 设置默认中文字体（Word 自带）
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal'].font.size = Pt(11)

        # 对 East Asian 脚本显式指定中文字体（python-docx 的中英混合字体设置）
        try:
            from docx.oxml.ns import qn
            doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except Exception:
            pass

        # 设置上下左右页边距
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        def _add_inline_runs(paragraph, element):
            """把 element 中的所有 inline 文本、加粗斜体等添加到 paragraph。"""
            for child in element.children:
                text = str(child)
                if isinstance(child, str):
                    # 纯文本直接添加
                    if text.strip() or text:
                        run = paragraph.add_run(text)
                        run.font.name = '微软雅黑'
                        try:
                            from docx.oxml.ns import qn
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                from docx.oxml import OxmlElement
                                rFonts = OxmlElement('w:rFonts')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:eastAsia'), '微软雅黑')
                            rFonts.set(qn('w:ascii'), '微软雅黑')
                            rFonts.set(qn('w:hAnsi'), '微软雅黑')
                        except Exception:
                            pass
                else:
                    tag = child.name.lower()
                    if tag in ('strong', 'b'):
                        text = child.get_text()
                        if text:
                            run = paragraph.add_run(text)
                            run.bold = True
                    elif tag in ('em', 'i'):
                        text = child.get_text()
                        if text:
                            run = paragraph.add_run(text)
                            run.italic = True
                    elif tag == 'code':
                        text = child.get_text()
                        if text:
                            run = paragraph.add_run(text)
                            run.font.name = 'Consolas'
                    elif tag in ('span', 'a'):
                        _add_inline_runs(paragraph, child)
                    else:
                        text = child.get_text()
                        if text:
                            paragraph.add_run(text)

        def _render_elem(elem):
            tag = elem.name and elem.name.lower()
            if tag is None:
                return
            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(tag[1])
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
                _add_inline_runs(para, elem)
                for run in para.runs:
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(18)
                    elif level == 2:
                        run.font.size = Pt(15)
                    else:
                        run.font.size = Pt(13)
                para.space_before = Pt(12)
                para.space_after = Pt(6)
            elif tag == 'p':
                para = doc.add_paragraph()
                _add_inline_runs(para, elem)
            elif tag in ('ul', 'ol'):
                ordered = (tag == 'ol')
                for i, li in enumerate(elem.find_all('li', recursive=False), start=1):
                    para = doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
                    _add_inline_runs(para, li)
            elif tag == 'table':
                rows = elem.find_all('tr')
                if not rows:
                    return
                n_cols = max(len(r.find_all(['th', 'td'])) for r in rows)
                n_cols = n_cols or 1
                table = doc.add_table(rows=len(rows), cols=n_cols)
                for r_idx, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for c_idx, cell in enumerate(cells):
                        if c_idx >= n_cols:
                            continue
                        docx_cell = table.rows[r_idx].cells[c_idx]
                        text = cell.get_text(strip=True)
                        docx_cell.text = text
                        is_header = (cell.name.lower() == 'th') or (r_idx == 0)
                        for para in docx_cell.paragraphs:
                            for run in para.runs:
                                if is_header:
                                    run.bold = True
                try:
                    table.style = 'Table Grid'
                except Exception:
                    pass
            elif tag == 'pre':
                para = doc.add_paragraph()
                for line in elem.get_text().splitlines():
                    run = para.add_run(line + '\n')
                    run.font.name = 'Consolas'
            elif tag == 'hr':
                doc.add_paragraph('_' * 40)
            elif tag == 'blockquote':
                para = doc.add_paragraph()
                _add_inline_runs(para, elem)
                for run in para.runs:
                    run.italic = True
            elif tag in ('div', 'section', 'article', 'html', 'body'):
                for child in elem.children:
                    if hasattr(child, 'name') and child.name:
                        _render_elem(child)
                    else:
                        t = str(child).strip()
                        if t:
                            doc.add_paragraph(t)
            else:
                t = elem.get_text(strip=True)
                if t:
                    doc.add_paragraph(t)

        body = soup.body or soup
        for child in body.children:
            if hasattr(child, 'name') and child.name:
                _render_elem(child)
            else:
                t = str(child).strip()
                if t:
                    doc.add_paragraph(t)

        doc.save(output_path)
        return True

    def _export(self):
        content = self.output_text.get("1.0", tk.END).strip()
        if not content:
            messagebox.showwarning("提示", "没有可导出的内容。")
            return

        if HAS_DOCX_SUPPORT:
            filetypes = [("Word文件", "*.docx"), ("文本文件", "*.txt")]
            default_ext = ".docx"
        else:
            filetypes = [("文本文件", "*.txt")]
            default_ext = ".txt"

        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile=f"{self.current_doc}_{self._timestamp()}"
        )
        if not file_path:
            return

        if file_path.lower().endswith('.docx'):
            if not HAS_DOCX_SUPPORT:
                messagebox.showwarning("提示", "Word导出功能不可用，请安装 python-docx 和 beautifulsoup4 库。")
                return
            try:
                self._markdown_to_docx(content, file_path)
                self.status_var.set("✅ 已导出Word")
                self.status_label.configure(fg=COLORS["success_green"])
            except Exception as e:
                messagebox.showerror("错误", f"Word导出失败：{str(e)}")
        else:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            self.status_var.set("✅ 已导出")
            self.status_label.configure(fg=COLORS["success_green"])

    def _copy_result(self):
        content = self.output_text.get("1.0", tk.END).strip()
        if content:
            self.root.clipboard_clear()
            self.root.clipboard_append(content)
            self.status_var.set("✅ 已复制")
            self.status_label.configure(fg=COLORS["success_green"])

    def _timestamp(self):
        return datetime.now().strftime("%Y%m%d_%H%M%S")

    def _show_settings(self):
        self._clear_content()
        self._update_nav_active(None)

        header = tk.Frame(self.content_frame, bg=COLORS["bg_card"])
        header.pack(fill=tk.X, padx=20, pady=(15, 0))
        tk.Label(header, text="⚙  系统设置", font=(FONT_FAMILY, 16, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=15, pady=12)

        body = tk.Frame(self.content_frame, bg=COLORS["bg_content"])
        body.pack(fill=tk.BOTH, expand=True, padx=20, pady=15)
        body.columnconfigure(0, weight=1)
        body.rowconfigure(1, weight=1)

        api_card = tk.Frame(body, bg=COLORS["bg_card"],
                            highlightbackground=COLORS["border_light"],
                            highlightthickness=1)
        api_card.grid(row=0, column=0, sticky="ew", pady=(0, 10))

        api_header = tk.Frame(api_card, bg=COLORS["bg_card"])
        api_header.pack(fill=tk.X)
        tk.Label(api_header, text="  API 配置", font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=10, pady=8)
        tk.Frame(api_header, bg=COLORS["btn_primary"], height=2).pack(fill=tk.X, padx=10)

        fields_frame = tk.Frame(api_card, bg=COLORS["bg_card"])
        fields_frame.pack(fill=tk.X, padx=20, pady=15)
        fields_frame.columnconfigure(1, weight=1)

        fields = [
            ("API 地址 (URL):", "url"),
            ("API Key:", "api_key"),
            ("模型名称:", "model"),
            ("Temperature:", "temperature"),
        ]

        self.setting_vars = {}
        for i, (label, key) in enumerate(fields):
            tk.Label(fields_frame, text=label, font=(FONT_FAMILY, 9),
                     bg=COLORS["bg_card"], fg=COLORS["text_secondary"]).grid(
                row=i, column=0, sticky="w", padx=(0, 15), pady=6)
            var = tk.StringVar(value=str(self.config.get(key, "")))
            entry = tk.Entry(fields_frame, textvariable=var, width=50,
                             font=(FONT_FAMILY_MONO, 9),
                             bg=COLORS["bg_input"], fg=COLORS["text_primary"],
                             insertbackground=COLORS["text_primary"],
                             relief=tk.FLAT, borderwidth=1,
                             highlightbackground=COLORS["border_light"],
                             highlightthickness=1)
            entry.grid(row=i, column=1, sticky="ew", pady=6)
            self.setting_vars[key] = var

        btn_row = tk.Frame(api_card, bg=COLORS["bg_card"])
        btn_row.pack(fill=tk.X, padx=20, pady=(0, 15))

        FlatButton(btn_row, text="保存配置", icon="💾",
                   command=self._save_settings,
                   color=COLORS["btn_primary"],
                   hover_color=COLORS["btn_primary_hover"]).pack(side=tk.LEFT, padx=(0, 10))
        FlatButton(btn_row, text="返回首页", icon="🏠",
                   command=self._show_home,
                   color="#7F8C8D", hover_color="#95A5A6").pack(side=tk.LEFT)

        prompt_card = tk.Frame(body, bg=COLORS["bg_card"],
                               highlightbackground=COLORS["border_light"],
                               highlightthickness=1)
        prompt_card.grid(row=1, column=0, sticky="nsew")

        prompt_header = tk.Frame(prompt_card, bg=COLORS["bg_card"])
        prompt_header.pack(fill=tk.X)
        tk.Label(prompt_header, text="  提示词模板管理", font=(FONT_FAMILY, 10, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=10, pady=8)
        tk.Frame(prompt_header, bg=COLORS["accent_gold"], height=2).pack(fill=tk.X, padx=10)

        tk.Label(prompt_card, text=f"模板目录：{os.path.join(BASE_DIR, 'prompt')}",
                 font=(FONT_FAMILY_MONO, 8), bg=COLORS["bg_card"],
                 fg=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(8, 2))
        tk.Label(prompt_card, text="为每种文书创建对应的 .txt 文件即可自定义提示词，文件名需与文书名称一致",
                 font=(FONT_FAMILY, 8), bg=COLORS["bg_card"],
                 fg=COLORS["text_secondary"]).pack(anchor="w", padx=15, pady=(0, 5))

        list_frame = tk.Frame(prompt_card, bg=COLORS["bg_card"])
        list_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)

        self.prompt_list = tk.Listbox(
            list_frame, font=(FONT_FAMILY_MONO, 9),
            bg=COLORS["bg_input"], fg=COLORS["text_primary"],
            selectbackground=COLORS["btn_primary"],
            selectforeground="white",
            relief=tk.FLAT, borderwidth=1,
            highlightbackground=COLORS["border_light"],
            highlightthickness=1, height=6
        )
        self.prompt_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        list_scroll = tk.Scrollbar(list_frame, command=self.prompt_list.yview)
        list_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.prompt_list.config(yscrollcommand=list_scroll.set)

        self._refresh_prompt_list()

        prompt_btn_row = tk.Frame(prompt_card, bg=COLORS["bg_card"])
        prompt_btn_row.pack(fill=tk.X, padx=15, pady=10)
        FlatButton(prompt_btn_row, text="编辑模板", icon="✏️",
                   command=self._edit_prompt,
                   color=COLORS["btn_primary"],
                   hover_color=COLORS["btn_primary_hover"]).pack(side=tk.LEFT, padx=(0, 8))
        FlatButton(prompt_btn_row, text="刷新列表", icon="🔄",
                   command=self._refresh_prompt_list,
                   color="#7F8C8D", hover_color="#95A5A6").pack(side=tk.LEFT)

    def _refresh_prompt_list(self):
        self.prompt_list.delete(0, tk.END)
        seen = set()
        for search_dir in [os.path.join(BASE_DIR, "prompt"), os.path.join(RESOURCE_DIR, "prompt")]:
            if os.path.exists(search_dir):
                for f in sorted(os.listdir(search_dir)):
                    if f.endswith(".txt") and f not in seen:
                        seen.add(f)
                        self.prompt_list.insert(tk.END, f)

    def _edit_prompt(self):
        sel = self.prompt_list.curselection()
        if not sel:
            messagebox.showwarning("提示", "请先选择一个模板文件。")
            return
        filename = self.prompt_list.get(sel[0])
        filepath = _find_prompt_file(filename.replace(".txt", "")) or os.path.join(RESOURCE_DIR, "prompt", filename)

        edit_win = tk.Toplevel(self.root)
        edit_win.title(f"编辑提示词 - {filename}")
        edit_win.geometry("650x450")
        edit_win.configure(bg=COLORS["bg_content"])

        header = tk.Frame(edit_win, bg=COLORS["bg_card"])
        header.pack(fill=tk.X)
        tk.Label(header, text=f"  ✏️  {filename}", font=(FONT_FAMILY, 12, "bold"),
                 bg=COLORS["bg_card"], fg=COLORS["text_primary"]).pack(side=tk.LEFT, padx=15, pady=10)

        text = scrolledtext.ScrolledText(
            edit_win, wrap=tk.WORD, font=(FONT_FAMILY, 10),
            bg=COLORS["bg_input"], fg=COLORS["text_primary"],
            insertbackground=COLORS["text_primary"],
            relief=tk.FLAT, borderwidth=0,
            selectbackground=COLORS["btn_primary"]
        )
        text.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        with open(filepath, "r", encoding="utf-8") as f:
            text.insert("1.0", f.read())

        def save():
            save_dir = os.path.join(BASE_DIR, "prompt")
            os.makedirs(save_dir, exist_ok=True)
            save_path = os.path.join(save_dir, filename)
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(text.get("1.0", tk.END))
            edit_win.destroy()
            self.status_var.set("✅ 模板已保存")

        btn_frame = tk.Frame(edit_win, bg=COLORS["bg_content"])
        btn_frame.pack(fill=tk.X, padx=15, pady=(0, 10))
        FlatButton(btn_frame, text="保存模板", icon="💾", command=save,
                   color=COLORS["btn_primary"],
                   hover_color=COLORS["btn_primary_hover"]).pack(side=tk.LEFT)
        FlatButton(btn_frame, text="取消", command=edit_win.destroy,
                   color="#7F8C8D", hover_color="#95A5A6").pack(side=tk.LEFT, padx=8)

    def _save_settings(self):
        config = configparser.ConfigParser()
        config["AI"] = {
            "URL": self.setting_vars["url"].get(),
            "API_KEY": self.setting_vars["api_key"].get(),
            "MODEL": self.setting_vars["model"].get(),
            "TEMPERATURE": self.setting_vars["temperature"].get()
        }
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            config.write(f)

        self.config = load_config()
        self.status_var.set("✅ 配置已保存")
        self.status_label.configure(fg=COLORS["success_green"])


def main():
    root = tk.Tk()

    style = ttk.Style()
    style.theme_use("clam")
    style.configure("TProgressbar", troughcolor=COLORS["border_light"],
                    background=COLORS["btn_primary"],
                    darkcolor=COLORS["btn_primary"],
                    lightcolor=COLORS["btn_primary_hover"],
                    bordercolor=COLORS["border_light"])

    app = LegalDocumentApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()